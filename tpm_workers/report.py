"""
tpm_workers.report - Maintenance Report Worker
ref: MASTER_PLAN_v5.md § 11.2.A

Pipeline (Researcher -> Writer -> Reviewer -> Maker):
  1. Researcher: load CM events + PM schedule + FMEA for target equipment
  2. Writer:     LLM generates narrative summary (Thai/English bilingual)
  3. Reviewer:   checklist (every claim has source row? no AI math? safety flag?)
  4. Maker:      python-docx renders to .docx
"""
from __future__ import annotations

import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any

from tpm_core.llm import chat
from tpm_workers.base import (
    WorkerInput,
    WorkerResult,
    WorkerStep,
    WorkerType,
)
from tpm_workers.data_loader import (
    CMHistoryLoader,
    FMEALoader,
    PMScheduleLoader,
    is_using_dummy_data,
)
from tpm_workers.metrics import full_summary

log = logging.getLogger(__name__)

DEFAULT_MODEL = os.getenv("TPM_WRITER_MODEL", "qwen3:8b")


# ============================================================
# Step 1: Researcher
# ============================================================
def step_researcher(inp: WorkerInput) -> tuple[WorkerStep, dict[str, Any]]:
    step = WorkerStep(name="researcher")
    cm = CMHistoryLoader()
    pm = PMScheduleLoader()
    fmea = FMEALoader()

    cm_df = cm.for_equipment(inp.target_subject)
    cm_recent = cm_df.copy()
    if not cm_df.empty and "Date" in cm_df.columns and inp.time_range_days:
        from datetime import timedelta
        cutoff = cm_df["Date"].max() - timedelta(days=inp.time_range_days)
        cm_recent = cm_df[cm_df["Date"] >= cutoff]

    pm_df = pm.for_equipment(inp.target_subject)
    fmea_df = fmea.for_equipment(inp.target_subject)

    summary = full_summary(cm_recent, observation_period_days=inp.time_range_days)

    payload = {
        "equipment_tag": inp.target_subject,
        "cm_total_count": int(len(cm_df)),
        "cm_recent_count": int(len(cm_recent)),
        "pm_tasks_count": int(len(pm_df)),
        "fmea_count": int(len(fmea_df)),
        "metrics": summary,
        # Sample raw rows for the writer to cite (top 5 by date)
        "recent_events_sample": cm_recent.head(5).to_dict("records"),
        "top_fmea": fmea_df.head(3).to_dict("records") if not fmea_df.empty else [],
        "pm_tasks": pm_df.to_dict("records") if not pm_df.empty else [],
        "data_source": str(cm.path) if cm.path else "(none)",
        "is_dummy": is_using_dummy_data(),
    }
    step.output = payload
    if not cm_df.empty:
        step.notes.append(
            f"loaded {len(cm_df)} CM rows for {inp.target_subject!r} "
            f"({len(cm_recent)} in last {inp.time_range_days}d)"
        )
    else:
        step.notes.append(f"no CM data for {inp.target_subject!r}")
    step.finish()
    return step, payload


# ============================================================
# Step 2: Writer (LLM)
# ============================================================
WRITER_SYSTEM = """\
You are a TPM (Total Productive Maintenance) engineer writing a concise
maintenance report. Write in mixed Thai/English (mostly Thai) matching
the user's request language.

Hard rules:
  - NEVER invent numbers or dates. Only use values from the provided data.
  - Every claim that uses a number must reference the source field.
  - If data is missing, write "ไม่มีข้อมูล" - do NOT guess.
  - Use a neutral engineering tone (no marketing, no superlatives).
  - Highlight Safety/LOTO concerns explicitly when relevant.

Output structure (Markdown):
  # Maintenance Report: <equipment_tag>

  ## 1. Executive Summary
  (2-3 sentences in Thai)

  ## 2. Reliability Metrics (last <N> days)
  (table or bullets - cite the metric values)

  ## 3. Recent Events
  (top 3-5 events with date + problem + severity)

  ## 4. Top Failure Modes (Pareto)
  (top 3 from the data)

  ## 5. Recommended Actions
  (concrete, conditional: "If X then do Y because Z")

  ## 6. Confidence & Caveats
  (if data is dummy or sparse, note it)
"""


def step_writer(researcher_payload: dict[str, Any], model: str) -> tuple[WorkerStep, str]:
    step = WorkerStep(name="writer")
    # Compact research payload for prompt (avoid token blow-up)
    prompt_data = {
        "equipment_tag": researcher_payload["equipment_tag"],
        "n_events_recent": researcher_payload["cm_recent_count"],
        "n_events_total": researcher_payload["cm_total_count"],
        "metrics": researcher_payload["metrics"],
        "recent_events_sample": researcher_payload["recent_events_sample"][:5],
        "top_fmea": researcher_payload["top_fmea"][:3],
        "pm_tasks": researcher_payload["pm_tasks"][:5],
        "is_dummy_data": researcher_payload["is_dummy"],
    }

    user_msg = (
        "Write the maintenance report based on the following structured data.\n"
        "Return only the Markdown report, no extra commentary.\n\n"
        f"```json\n{json.dumps(prompt_data, ensure_ascii=False, indent=2, default=str)}\n```"
    )
    try:
        report_md = chat(
            model,
            [
                {"role": "system", "content": WRITER_SYSTEM},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.3,
            timeout=180.0,
        )
    except Exception as e:  # noqa: BLE001
        step.finish(success=False, error=f"writer failed: {e}")
        return step, ""
    step.output = {"report_md_chars": len(report_md), "model": model}
    step.notes.append(f"generated {len(report_md)} chars")
    step.finish()
    return step, report_md


# ============================================================
# Step 3: Reviewer (rule-based checklist)
# ============================================================
REQUIRED_SECTIONS = [
    "Executive Summary",
    "Reliability Metrics",
    "Recent Events",
    "Failure Modes",
    "Recommended Actions",
]


def step_reviewer(report_md: str, researcher_payload: dict[str, Any]) -> tuple[WorkerStep, list[str]]:
    step = WorkerStep(name="reviewer")
    findings: list[str] = []

    # Check 1: required sections present
    md_lower = report_md.lower()
    for section in REQUIRED_SECTIONS:
        if section.lower() not in md_lower:
            findings.append(f"missing section: {section}")

    # Check 2: at least one number from metrics appears in report
    metrics = researcher_payload.get("metrics", {})
    n_events = metrics.get("n_events", 0)
    if n_events > 0 and str(n_events) not in report_md:
        findings.append(
            f"metric n_events={n_events} not cited in report - possible hallucination"
        )

    # Check 3: dummy data should be flagged
    if researcher_payload.get("is_dummy") and "dummy" not in md_lower and "ทดสอบ" not in report_md:
        findings.append("dummy data not flagged in report (recommend adding caveat)")

    # Check 4: report is not empty / not too short
    if len(report_md) < 200:
        findings.append(f"report too short ({len(report_md)} chars) - likely failed")

    step.output = {
        "n_findings": len(findings),
        "findings": findings,
        "report_chars": len(report_md),
    }
    step.notes = findings or ["all checks passed"]
    # Reviewer fails the worker only on critical issues
    critical = [f for f in findings if "too short" in f or "missing section" in f]
    step.finish(success=len(critical) == 0)
    return step, findings


# ============================================================
# Step 4: Maker (.docx)
# ============================================================
def step_maker(
    report_md: str,
    researcher_payload: dict[str, Any],
    output_dir: Path,
    session_id: str,
) -> tuple[WorkerStep, list[str]]:
    step = WorkerStep(name="maker")
    output_dir.mkdir(parents=True, exist_ok=True)
    tag = (researcher_payload.get("equipment_tag") or "unknown").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"maintenance_report_{tag}_{timestamp}.docx"
    md_path = output_dir / f"maintenance_report_{tag}_{timestamp}.md"

    # Save raw markdown alongside (always - for inspection)
    md_path.write_text(report_md, encoding="utf-8")

    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        title = doc.add_heading(f"Maintenance Report: {researcher_payload.get('equipment_tag','')}", level=0)
        for run in title.runs:
            run.font.size = Pt(18)

        # Lineage block (per AGENTS.md § 7)
        meta = doc.add_paragraph()
        meta.add_run(f"Generated: {datetime.now().isoformat(timespec='seconds')}\n").italic = True
        meta.add_run(f"Session: {session_id}\n").italic = True
        meta.add_run(f"Source: {researcher_payload.get('data_source','?')}\n").italic = True
        if researcher_payload.get("is_dummy"):
            warn = doc.add_paragraph()
            r = warn.add_run("WARNING: This report uses DUMMY data (pre-internship testing).")
            r.bold = True

        # Render Markdown lightly (heading + paragraph; no full markdown parser)
        for line in report_md.splitlines():
            stripped = line.rstrip()
            if not stripped:
                doc.add_paragraph("")
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

        doc.save(docx_path)
        step.output = {"docx": str(docx_path), "md": str(md_path)}
        step.notes.append(f"wrote {docx_path.name}")
    except Exception as e:  # noqa: BLE001
        step.finish(success=False, error=f"docx render failed: {e}")
        return step, [str(md_path)]  # at least the .md was saved
    step.finish()
    return step, [str(docx_path), str(md_path)]


# ============================================================
# Top-level
# ============================================================
def run_report_worker(
    inp: WorkerInput,
    model: str = DEFAULT_MODEL,
) -> WorkerResult:
    result = WorkerResult(worker_type=WorkerType.REPORT)

    # 1. Researcher
    s1, payload = step_researcher(inp)
    result.add_step(s1)
    if payload["cm_total_count"] == 0:
        result.summary = (
            f"No CM history found for {inp.target_subject!r}. "
            f"Available equipment: see data_loader.list_equipment_tags()"
        )
        result.success = False
        return result

    # 2. Writer
    s2, report_md = step_writer(payload, model=model)
    result.add_step(s2)
    if not s2.success or not report_md:
        result.summary = f"Writer failed: {s2.error}"
        result.success = False
        return result

    # 3. Reviewer
    s3, findings = step_reviewer(report_md, payload)
    result.add_step(s3)
    result.auditor_findings = findings
    result.auditor_passed = s3.success

    # 4. Maker
    s4, files = step_maker(report_md, payload, inp.output_dir, inp.session_id)
    result.add_step(s4)
    result.output_files = files

    result.metrics = payload["metrics"]
    result.summary = (
        f"Report for {inp.target_subject}: "
        f"{payload['cm_recent_count']} events in last {inp.time_range_days}d, "
        f"MTBF={payload['metrics']['mtbf'].get('mtbf_hours', 'n/a')}h, "
        f"availability={payload['metrics']['availability_pct']}%"
    )
    result.confidence = 0.85 if s3.success else 0.55
    result.success = s4.success and bool(files)
    return result
